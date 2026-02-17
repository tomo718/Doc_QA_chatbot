# ===============================
# ドキュメント読込処理
# ===============================
import tempfile
from langchain_core.documents import Document
from docx import Document as DocxDocument
import pandas as pd


def load_documents_from_files(uploaded_files):
    documents = []

    for uploaded_file in uploaded_files:

        # 一時ファイルに保存
        with tempfile.NamedTemporaryFile(delete=False) as tmp:
            tmp.write(uploaded_file.read())
            tmp_path = tmp.name

        # ------------------------
        # 📄 Word処理
        # ------------------------
        if uploaded_file.name.endswith(".docx"):
            doc = DocxDocument(tmp_path)

            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            text = "\n".join(full_text)

            documents.append(
                Document(
                    page_content=text,
                    metadata={"source": uploaded_file.name}
                )
            )

        # ------------------------
        # 📊 Excel処理（全シート読み込み）
        # ------------------------
        elif uploaded_file.name.endswith(".xlsx"):
            xls = pd.ExcelFile(tmp_path)

            for sheet_name in xls.sheet_names:
                df = xls.parse(sheet_name)

                # NaNを空文字へ
                df = df.fillna("")

                # 表を文字列化
                text = f"【Sheet: {sheet_name}】\n"
                text += df.to_string(index=False)

                documents.append(
                    Document(
                        page_content=text,
                        metadata={"source": f"{uploaded_file.name} - {sheet_name}"}
                    )
                )

    return documents
